"""Test: python-docx"""
import sys
sys.stdout.reconfigure(line_buffering=True)
try:
    import lxml  # noqa: F401 — hard dependency
except ImportError:
    print("python-docx: SKIP (lxml not available)")
    sys.exit(0)
try:
    from docx import Document

    # Create document
    doc = Document()
    doc.add_heading("Nanvix Test", level=1)
    doc.add_paragraph("Hello from Nanvix!")

    assert len(doc.paragraphs) >= 1
    assert "Hello from Nanvix!" in doc.paragraphs[-1].text

    print("python-docx: PASS")
except Exception as e:
    print(f"python-docx: FAIL: {e}")
    sys.exit(1)
